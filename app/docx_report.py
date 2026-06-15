"""
docx_report.py — builds the Facility Assessment Snapshot as an editable Word
document (.docx), mirroring the report table. python-docx generates the file
in-process, the same way WeasyPrint produces the PDF.
"""

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOGO_PATH = Path(__file__).parent / "static" / "medelite-logo.png"

GREY  = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x15, 0x80, 0x3D)
RED   = RGBColor(0xB9, 0x1C, 0x1C)
LABEL_FILL = "F1F5F9"
GOOD_FILL  = "F0FDF4"
BAD_FILL   = "FEF2F2"


def _shade(cell, fill_hex):
    """Set a table cell's background (clear shading, never solid)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _add_hyperlink(paragraph, url, text):
    """Append a real clickable hyperlink to a paragraph."""
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "2563EB"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    paragraph._p.append(link)


def build_docx(data) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    # Banner — logo, or text fallback. Reuse the default first paragraph.
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        head.add_run().add_picture(str(LOGO_PATH), width=Inches(2.4))
    else:
        r = head.add_run("INFINITE — Managed by MEDELITE")
        r.bold = True; r.font.size = Pt(16)

    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("FACILITY ASSESSMENT SNAPSHOT"); tr.bold = True; tr.font.size = Pt(13)

    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(data["state"] or ""); sr.font.size = Pt(10); sr.font.color.rgb = GREY

    # rows: (label, value, fallback, tone)
    rows = [
        ("Name of Facility", data["name"], "Not available", None),
        ("Location", data["location"], "Not available", None),
        ("EMR", data["emr"], "Not provided", None),
        ("Census Capacity", data["certified_beds"], "Not available", None),
        ("Current Census", data["current_census"], "Not provided", None),
        ("Type of Patient", data["patient_type"], "Not provided", None),
        ("Previous Coverage from Medelite", data["previous_coverage"], "Not provided", None),
        ("Previous Provider Performance from Medelite", data["previous_performance"], "Not provided", None),
        ("Medical Coverage", data["medical_coverage"], "Not provided", None),
        ("Overall Star Rating", data["ratings"]["overall"], "Not available", None),
        ("Health Inspection", data["ratings"]["health_inspection"], "Not available", None),
        ("Staffing", data["ratings"]["staffing"], "Not available", None),
        ("Quality of Resident Care", data["ratings"]["quality_of_care"], "Not available", None),
    ]
    for label, value, tone in data["metrics_rows"]:
        rows.append((label, value, "Not available", tone))

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, (label, value, fallback, tone) in enumerate(rows):
        lcell, vcell = table.rows[i].cells
        lcell.width, vcell.width = Inches(2.9), Inches(3.6)

        _shade(lcell, LABEL_FILL)
        lr = lcell.paragraphs[0].add_run(label); lr.bold = True; lr.font.size = Pt(10)

        vp = vcell.paragraphs[0]
        text = str(value) if (value is not None and value != "") else fallback
        vp.add_run(text).font.size = Pt(10)

        if tone == "good":
            _shade(vcell, GOOD_FILL)
            b = vp.add_run("   ▼ below avg"); b.bold = True; b.font.size = Pt(8); b.font.color.rgb = GREEN
        elif tone == "bad":
            _shade(vcell, BAD_FILL)
            b = vp.add_run("   ▲ above avg"); b.bold = True; b.font.size = Pt(8); b.font.color.rgb = RED

    legend = doc.add_paragraph()
    g = legend.add_run("▼ below avg"); g.bold = True; g.font.size = Pt(8); g.font.color.rgb = GREEN
    g2 = legend.add_run(" = better than the national average; lower is better for these measures.")
    g2.font.size = Pt(8); g2.font.color.rgb = GREY

    footer = doc.add_paragraph()
    f = footer.add_run(f"Data as of {data['processing_date']} · ")
    f.font.size = Pt(8); f.font.color.rgb = GREY
    _add_hyperlink(footer, data["care_compare_url"], "View on Medicare Care Compare")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()